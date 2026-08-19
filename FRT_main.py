import numpy as np
#import cv2
import matplotlib.pyplot as plt
import os
import ipywidgets as widgets
from IPython.display import display
from scipy.optimize import curve_fit
from scipy.ndimage import zoom
from PIL import Image
from scipy.signal import fftconvolve
from scipy.fft import fft2, ifft2, fftshift
from scipy.ndimage import gaussian_filter
from scipy.signal import fftconvolve
from scipy.optimize import minimize
import numpy as np
import matplotlib.pyplot as plt
from scipy.signal import convolve2d as conv2
from skimage import color, data, restoration
from scipy.optimize import curve_fit
from skimage.restoration import richardson_lucy
from scipy.signal import convolve2d
from scipy.ndimage import convolve
from mpl_toolkits.mplot3d import Axes3D
from IPython.display import display, HTML
import pandas as pd
import ipywidgets as widgets
import plotly.graph_objects as go
from IPython.display import display
import ipywidgets as widgets
from scipy.interpolate import griddata
import numpy as np
from scipy.optimize import minimize
from scipy.special import erf
from tkinter import Tk, filedialog
import ipywidgets as widgets
from IPython.display import display
import io
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os


def normalize_data(pixels):
    """Нормировка яркости так, чтобы сумма была равна 1"""
    pixels = np.array(pixels, dtype=float)
    return pixels / pixels.sum()


def gaussian_pixel_integral(x0, y0, sigma, i, j):
    """
    Интеграл гауссианы по площади пикселя (i,j).
    Считаем, что пиксели имеют координаты с центрами в целых числах,
    а их границы: [i-0.5, i+0.5], [j-0.5, j+0.5].
    """

    def phi(t):
        return 0.5 * (1 + erf(t / (np.sqrt(2) * sigma)))

    x1, x2 = i - 0.5, i + 0.5
    y1, y2 = j - 0.5, j + 0.5

    return (phi(x2 - x0) - phi(x1 - x0)) * (phi(y2 - y0) - phi(y1 - y0))


def model_image(shape, x0, y0, sigma):
    """Модельное распределение для матрицы пикселей shape=(ny,nx)"""
    ny, nx = shape
    g = np.zeros((ny, nx))
    for j in range(ny):
        for i in range(nx):
            g[j, i] = gaussian_pixel_integral(x0, y0, sigma, i, j)
    return g / g.sum()  # нормировка на 1


def fit_gaussian(pixels):
    """Подгонка параметров (x0, y0, sigma) под данные"""
    data = normalize_data(pixels)
    ny, nx = data.shape

    # Начальные приближения
    x0_init, y0_init = nx / 2, ny / 2
    sigma_init = 1.0

    def loss(params):
        x0, y0, sigma = params
        if sigma <= 0:
            return np.inf
        model = model_image((ny, nx), x0, y0, sigma)
        return np.sum((data - model) ** 2)

    res = minimize(loss, [x0_init, y0_init, sigma_init],
                   method="Nelder-Mead")
    x0, y0, sigma = res.x
    return {"x0": x0, "y0": y0, "sigma": sigma, "success": res.success, "loss": res.fun}


def fit_gaussian_with_model(pixels):
    """
    Подгонка 2D гауссианы под матрицу пикселей.

    Возвращает:
    - params: x0, y0, sigma
    - model: матрица гауссианы по пикселям (нормализованная)
    """
    data = normalize_data(pixels)
    ny, nx = data.shape

    # Начальные приближения
    x0_init, y0_init = nx / 2, ny / 2
    sigma_init = 1.0

    def loss(params):
        x0, y0, sigma = params
        if sigma <= 0:
            return np.inf
        model = model_image((ny, nx), x0, y0, sigma)
        return np.sum((data - model) ** 2)

    # Оптимизация
    res = minimize(loss, [x0_init, y0_init, sigma_init], method="Nelder-Mead")
    x0, y0, sigma = res.x

    # Строим модель с оптимальными параметрами
    model = model_image((ny, nx), x0, y0, sigma)

    return {
        "x0": x0,
        "y0": y0,
        "sigma": sigma,
        "success": res.success,
        "loss": res.fun,
        "model": model
    }


def fit_gaussian_central(pixels):
    """
    Подгонка 2D гауссианы под матрицу пикселей.
    Центральная точка будет идеально совпадать по амплитуде.

    Возвращает:
    - params: x0, y0, sigma, A
    - model: матрица гауссианы по пикселям
    - abs_errors: матрица абсолютных ошибок
    """
    pixels = np.array(pixels, dtype=float)
    ny, nx = pixels.shape

    # Начальные приближения
    x0_init, y0_init = nx / 2, ny / 2
    sigma_init = 1.0

    def loss(params):
        """Функция ошибки для оптимизации (без амплитуды)"""
        x0, y0, sigma = params
        if sigma <= 0:
            return np.inf
        model = model_image((ny, nx), x0, y0, sigma)
        # Минимизируем суммарную ошибку по всем пикселям
        return np.sum((pixels - model) ** 2)

    # Оптимизация
    res = minimize(loss, [x0_init, y0_init, sigma_init], method="Nelder-Mead")
    x0, y0, sigma = res.x

    # Вычисляем амплитуду так, чтобы центральный пиксель совпадал
    yc, xc = ny // 2, nx // 2
    central_value = pixels[yc, xc]
    central_gauss = model_image((ny, nx), x0, y0, sigma)[yc, xc]
    A = central_value / central_gauss

    # Строим окончательную модель
    model = A * model_image((ny, nx), x0, y0, sigma)

    # Матрица абсолютных ошибок
    abs_errors = np.abs(pixels - model)

    return {
        "x0": x0,
        "y0": y0,
        "sigma": sigma,
        "A": A,
        "success": res.success,
        "loss": res.fun,
        "model": model,
        "abs_errors": abs_errors
    }


def fit_gaussian_weighted(pixels):
    """
    Подгонка 2D гауссианы под матрицу пикселей с учётом весов.
    Центральная точка будет идеально совпадать по амплитуде.
    Чем ярче пиксель, тем выше его вес.

    Возвращает:
    - params: x0, y0, sigma, A
    - model: матрица гауссианы по пикселям
    - abs_errors: матрица абсолютных ошибок
    """
    pixels = np.array(pixels, dtype=float)
    ny, nx = pixels.shape

    # Вес пикселя пропорционален его значению
    # Можно добавить небольшую константу, чтобы тёмные пиксели тоже учитывались
    W = pixels

    # Начальные приближения
    x0_init, y0_init = nx / 2, ny / 2
    sigma_init = 1.0

    def loss(params):
        x0, y0, sigma = params
        if sigma <= 0:
            return np.inf
        model = model_image((ny, nx), x0, y0, sigma)
        # Взвешенная функция ошибки
        return np.sum(W * (pixels - model) ** 2)

    # Оптимизация
    res = minimize(loss, [x0_init, y0_init, sigma_init], method="Nelder-Mead")
    x0, y0, sigma = res.x

    # Амплитуда для центрального пикселя
    yc, xc = ny // 2, nx // 2
    central_value = pixels[yc, xc]
    central_gauss = model_image((ny, nx), x0, y0, sigma)[yc, xc]
    A = central_value / central_gauss

    # Финальная модель
    model = A * model_image((ny, nx), x0, y0, sigma)

    # Абсолютная погрешность
    abs_errors = np.abs(pixels - model)

    return {
        "x0": x0,
        "y0": y0,
        "sigma": sigma,
        "A": A,
        "success": res.success,
        "loss": res.fun,
        "model": model,
        "abs_errors": abs_errors,
        "weights": W
    }


def normalize_pixels_sum1(pixels):
    """
    Нормализует матрицу пикселей так, чтобы:
    1. Все значения были в диапазоне [0, 1]
    2. Сумма всех элементов была равна 1
    """
    pixels = np.array(pixels, dtype=float)

    # сдвиг к нулю, если есть отрицательные
    pixels = pixels - pixels.min()

    # нормализация диапазона к [0, 1]
    max_val = pixels.max()
    if max_val > 0:
        pixels = pixels / max_val

    # нормализация суммы к 1
    total = pixels.sum()
    if total > 0:
        pixels = pixels / total

    return pixels


# ==== функции для оформления ====
def set_font(run, font_name="Times New Roman", size=12):
    run.font.name = font_name
    run.font.size = Pt(size)


def add_matrix_table(doc, matrix, title):
    """Добавляет таблицу с матрицей"""
    p = doc.add_paragraph(title)
    set_font(p.runs[0])
    table = doc.add_table(rows=matrix.shape[0], cols=matrix.shape[1])
    table.style = "Table Grid"
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            cell = table.cell(i, j)
            cell.text = f"{matrix[i, j]:.3f}"
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    set_font(run)


def add_params_table(doc, params_dict):
    """Добавляет таблицу параметров гауссианы"""
    p = doc.add_paragraph("Параметры гауссианы:")
    set_font(p.runs[0])
    table = doc.add_table(rows=len(params_dict), cols=2)
    table.style = "Table Grid"
    for i, (key, value) in enumerate(params_dict.items()):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = f"{value:.3f}" if isinstance(value, (int, float, np.floating)) else str(value)
        for j in range(2):
            for par in table.cell(i, j).paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    set_font(run)


def add_centered_image_with_caption(doc, img_path, caption_text, width=Inches(3.5)):
    """Вставляет изображение по центру с подписью"""
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(img_path, width=width)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # подпись
    p_caption = doc.add_paragraph(caption_text)
    for run in p_caption.runs:
        set_font(run)
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def E_h2(D, DI):
    """Освещенность на зрачке от дальности D (км)."""
    term1 = DI / ((D * 1e5) ** 2)
    term2 = 10 ** (-0.6 - D / 190 + np.exp(-((D + 43) / 47) ** 2.9)) - 0.04
    return term1 * term2

def snr_from_D(D, DI, K=1.0, sigma_n=1.0):
    # Чтобы не было отрицательных значений из-за '-0.04' в term2:
    E = np.maximum(E_h2(D, DI), 0.0)
    return K * E / sigma_n

def Dmax_from_snr_threshold(snr_th, DI, K=1.0, sigma_n=1.0,
                            D_min=0.1, D_max=120.0, n_grid=20000):
    """
    Максимальная дальность, где SNR(D) >= snr_th.
    Поиск по плотной сетке (надежно и просто).
    """
    D_grid = np.linspace(D_min, D_max, n_grid)
    snr_grid = snr_from_D(D_grid, DI=DI, K=K, sigma_n=sigma_n)

    idx = np.where(snr_grid >= snr_th)[0]
    if len(idx) == 0:
        return np.nan
    return D_grid[idx[-1]]  # самая дальняя точка, где порог еще выполняется

def plot_Dmax_vs_SNR(DI, K=1.0, sigma_n=1.0,
                     snr_min_db=-10, snr_max_db=30, n_points=120):
    snr_db = np.linspace(snr_min_db, snr_max_db, n_points)
    snr_lin = 10 ** (snr_db / 10.0)

    Dmax = np.array([
        Dmax_from_snr_threshold(s, DI=DI, K=K, sigma_n=sigma_n)
        for s in snr_lin
    ])

    plt.figure(figsize=(8,5))
    plt.plot(snr_db, Dmax, lw=2)
    plt.xlabel('Требуемый порог SNR, дБ')
    plt.ylabel('Максимальная дальность, км')
    plt.title('Dmax(SNR) по модели E_h2(D, DI)')
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.show()

# === Шаг 1. Загружаем матрицу из файла ===
filename = "C:/Users/Админ/Downloads/Test_image/Dump_Beo-M_MPO_MW_SRAM (1).txt"
data = np.loadtxt(filename)

print("Исходная матрица:", data.shape)

# === Шаг 2. Обрезаем по 5 строк и столбцов с каждой стороны ===
cropped = data[5:-5, 5:-5]
print("После обрезки:", cropped.shape)

# === Шаг 3. Находим пиксель с максимальной яркостью ===
y_max, x_max = np.unravel_index(np.argmax(cropped), cropped.shape)
print("Максимум:", cropped[y_max, x_max], "в точке (y,x) =", (y_max, x_max))

# === Шаг 4. Вырезаем окно 3x3 вокруг максимума ===
y_start, y_end = max(0, y_max - 1), min(cropped.shape[0], y_max + 2)
x_start, x_end = max(0, x_max - 1), min(cropped.shape[1], x_max + 2)
roi = cropped[y_start:y_end, x_start:x_end]

print("\nМатрица 3x3 вокруг ярчайшего пикселя:")
print(roi)

# === Функция для выборки квадратов ===
def get_ring(cropped, y_max, x_max, size):
    """Возвращает квадрат size×size вокруг (y_max, x_max)."""
    half = size // 2
    y0, y1 = y_max - half, y_max + half + 1
    x0, x1 = x_max - half, x_max + half + 1
    return cropped[y0:y1, x0:x1]

# квадрат 5x5
win5 = get_ring(cropped, y_max, x_max, 5)
ring5 = win5.copy()
ring5[1:-1, 1:-1] = np.nan   # убираем 3x3

# квадрат 9x9
win9 = get_ring(cropped, y_max, x_max, 9)
ring9 = win9.copy()
ring9[2:-2, 2:-2] = np.nan   # убираем 5x5

# собираем все значения, исключая nan
background_pixels = np.concatenate([ring5[~np.isnan(ring5)], ring9[~np.isnan(ring9)]])
background = np.mean(background_pixels)

print("\nФон (усреднённое значение 18+56 пикселей):", background)

# === Шаг 4.2. Вычитаем фон из 3x3 ===
roi_corrected = roi - background
print("\nМатрица 3x3 после вычитания фона:")
print(roi_corrected)

# === Шаг 5. Строим изображение именно для 3x3 ===
plt.imshow(roi_corrected, cmap="gray", interpolation="nearest")
plt.colorbar(label="Яркость (с вычетом фона)")
plt.title("Фрагмент 3x3 вокруг максимума (фон вычтен)")
plt.show()

pixels = normalize_pixels_sum1(roi_corrected)
result = fit_gaussian_weighted(pixels)

print("Параметры гауссианы:")
print("x0 =", result['x0'], "y0 =", result['y0'], "sigma =", result['sigma'], "A =", result['A'])

print("\nМатрица модели:")
print(result['model'])

print("\nМатрица абсолютных ошибок:")
print(result['abs_errors'])

print("\nМатрица весов:")
print(result['weights'])

# === Визуализация центра гауссианы с окружностями ===
plt.imshow(roi_corrected, cmap="gray", interpolation="nearest")
plt.colorbar(label="Яркость (с вычетом фона)")
plt.title("Фрагмент 3x3 с центром гауссианы")

# Красная точка - центр
plt.plot(result['x0'], result['y0'], 'ro', markersize=3, label="Центр гауссианы")

# Добавляем окружности
circle_params = [
    (result['sigma'], 'r', "σ"),
    (1.52 * result['sigma'], 'g', "1.52σ (≈80%)"),
    (1.73 * result['sigma'], 'b', "1.73σ (≈90%)"),
    (2.3  * result['sigma'], 'r', "2.3σ (≈99%)")
]

ax = plt.gca()
for radius, color, label in circle_params:
    circle = patches.Circle((result['x0'], result['y0']), radius,
                            edgecolor=color, facecolor='none',
                            linestyle='--', linewidth=1, label=label)
    ax.add_patch(circle)

plt.legend()
plt.show()